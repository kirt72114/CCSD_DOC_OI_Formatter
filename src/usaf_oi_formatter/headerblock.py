"""DAFMAN 90-161 Figure A2.2 title block builder.

Replaces whatever is above the first heading paragraph with a standards-
compliant title block. Uses python-docx's `add_table` so we get its
registered element classes, then moves the resulting XML into the
correct position within the body.
"""

from __future__ import annotations

from docx.document import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

from . import rules
from .meta import OIMeta
from .profile import FormattingProfile, default as _default_profile


def rebuild(doc: Document, meta: OIMeta,
            profile: FormattingProfile | None = None) -> None:
    profile = profile or _default_profile()
    # 1. Build the pieces at the end of the doc using python-docx APIs.
    top_table = _build_top_table(doc, meta, profile)
    compliance_p = _build_compliance_line(doc)
    access_p, release_p = _build_access_release(doc, meta, profile)
    opr_table = _build_opr_table(doc, meta, profile)
    rule_p = _build_horizontal_rule(doc)

    # 2. Collect their XML elements in order.
    head_elements = [
        top_table._tbl,
        compliance_p._p,
        access_p._p,
        release_p._p,
        opr_table._tbl,
        rule_p._p,
    ]

    # 3. Remove any prefix the original doc had; leave headings intact.
    _remove_existing_prefix(doc)

    # 4. Insert the title block at the top of the body.
    body = doc.element.body
    for i, elem in enumerate(head_elements):
        body.insert(i, elem)


def _remove_existing_prefix(doc: Document) -> None:
    body = doc.element.body
    first_heading = _find_first_heading(doc)
    if first_heading is None:
        return
    for child in list(body):
        if child is first_heading._p:
            break
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _find_first_heading(doc: Document):
    heading_names = set(rules.HEADING_STYLES) | {"Heading 1", "Heading 2", "Heading 3"}
    for p in doc.paragraphs:
        if p.style.name in heading_names:
            return p
    return None


# ---------- component builders ---------------------------------------

def _build_top_table(doc: Document, meta: OIMeta, profile: FormattingProfile):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for col in table.columns:
        col.width = Inches(3.25)
    _kill_borders(table)

    _set_cell_lines(
        doc, table.cell(0, 0), profile,
        [(rules.LBL_BYORDER, True),
         (_nonempty(meta.unit.upper(), "UNIT"), False)],
    )
    _set_cell_lines(
        doc, table.cell(0, 1), profile,
        [(_nonempty(meta.oi_number.upper(), "UNIT OPERATING INSTRUCTION XX-X"), True),
         (_nonempty(meta.date_str, "DD Month YYYY"), False),
         ("", False),
         (_nonempty(meta.category, "Category"), False),
         (_nonempty(meta.subject, "Subject"), False)],
    )
    return table


def _build_compliance_line(doc: Document):
    p = doc.add_paragraph()
    p.style = doc.styles[rules.STY_TITLEBLOCK]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(rules.LBL_COMPLIANCE)
    run.bold = True
    return p


def _build_access_release(doc: Document, meta: OIMeta, profile: FormattingProfile):
    acc = _nonempty(meta.accessibility, profile.default_accessibility)
    rel = _nonempty(meta.releasability, profile.default_releasability)

    access_p = doc.add_paragraph(f"{rules.LBL_ACCESSIBILITY}  {acc}")
    access_p.style = doc.styles[rules.STY_TITLEBLOCK]

    release_p = doc.add_paragraph(f"{rules.LBL_RELEASABILITY}  {rel}")
    release_p.style = doc.styles[rules.STY_TITLEBLOCK]
    return access_p, release_p


def _build_opr_table(doc: Document, meta: OIMeta, profile: FormattingProfile):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    for col in table.columns:
        col.width = Inches(3.25)
    _kill_borders(table)

    _set_cell_lines(
        doc, table.cell(0, 0), profile,
        [(f"{rules.LBL_OPR} {_nonempty(meta.opr, 'OPR')}", False)],
    )
    _set_cell_lines(
        doc, table.cell(0, 1), profile,
        [(f"{rules.LBL_CERTIFIED_BY} {_nonempty(meta.certified_by, 'TBD')}", False)],
    )
    _set_cell_lines(
        doc, table.cell(1, 0), profile,
        [(f"{rules.LBL_SUPERSEDES} {_nonempty(meta.supersedes, 'N/A')}", False)],
    )
    _set_cell_lines(
        doc, table.cell(1, 1), profile,
        [(f"{rules.LBL_PAGES} {_nonempty(meta.pages, 'TBD')}", False)],
    )
    return table


def _build_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.style = doc.styles[rules.STY_TITLEBLOCK]
    pPr = p._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    return p


# ---------- helpers --------------------------------------------------

def _kill_borders(table) -> None:
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = etree.SubElement(tblPr, qn("w:tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = etree.SubElement(borders, qn(f"w:{side}"))
        b.set(qn("w:val"), "nil")


def _set_cell_lines(doc: Document, cell, profile: FormattingProfile,
                    lines: list[tuple[str, bool]]) -> None:
    cell.text = ""
    first_para = cell.paragraphs[0]
    first_para.style = doc.styles[rules.STY_TITLEBLOCK]

    for i, (line, bold) in enumerate(lines):
        if i == 0:
            p = first_para
        else:
            p = cell.add_paragraph(style=doc.styles[rules.STY_TITLEBLOCK])
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        run = p.add_run(line)
        run.bold = bold
        run.font.name = profile.titleblock_font
        run.font.size = Pt(profile.titleblock_size_pt)


def _nonempty(value: str, fallback: str) -> str:
    return value.strip() or fallback
