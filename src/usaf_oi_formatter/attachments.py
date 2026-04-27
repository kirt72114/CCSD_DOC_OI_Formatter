"""Rebuild `Attachment N—TITLE` headings and seed a Glossary attachment."""

from __future__ import annotations

import re

from docx.document import Document

from . import rules
from .profile import FormattingProfile, default as _default_profile

_ATTACH_RE = re.compile(
    rf"^\s*attachment\s+(\d+)\s*[\-{rules.ATTACH_SEP}–:.]?\s*(.*)$",
    re.IGNORECASE,
)


def apply(doc: Document, glossary: dict[str, str],
          profile: FormattingProfile | None = None) -> None:
    p = profile or _default_profile()
    _normalize_existing(doc)
    if p.seed_glossary and not _has_attachment_1(doc):
        _insert_glossary(doc, glossary)


def _normalize_existing(doc: Document) -> None:
    for p in doc.paragraphs:
        match = _ATTACH_RE.match(p.text.strip())
        if not match:
            continue
        num, rest = match.group(1), match.group(2).strip().upper()
        new_text = f"{rules.ATTACH_PREFIX}{num}{rules.ATTACH_SEP}{rest}"
        _replace_paragraph_text(p, new_text)
        p.style = doc.styles[rules.STY_ATTACH_TITLE]


def _has_attachment_1(doc: Document) -> bool:
    prefix = f"{rules.ATTACH_PREFIX}1{rules.ATTACH_SEP}"
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return True
    return False


def _insert_glossary(doc: Document, glossary: dict[str, str]) -> None:
    title_text = (
        f"{rules.ATTACH_PREFIX}1{rules.ATTACH_SEP}{rules.GLOSSARY_TITLE}"
    )
    title_p = doc.add_paragraph(title_text)
    title_p.style = doc.styles[rules.STY_ATTACH_TITLE]

    if not glossary:
        return

    for acronym in sorted(glossary):
        expansion = glossary[acronym]
        entry = doc.add_paragraph(f"{acronym}\t{expansion}")
        entry.style = doc.styles[rules.STY_BODY]


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(new_text)
