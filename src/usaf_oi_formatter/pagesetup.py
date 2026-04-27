"""Page setup: margins, paper size, orientation, page numbers."""

from __future__ import annotations

from docx.document import Document
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.shared import Inches

from .profile import FormattingProfile, default as _default_profile


def apply(doc: Document, profile: FormattingProfile | None = None) -> None:
    p = profile or _default_profile()
    for section in doc.sections:
        section.top_margin = Inches(p.margin_top_in)
        section.bottom_margin = Inches(p.margin_bottom_in)
        section.left_margin = Inches(p.margin_left_in)
        section.right_margin = Inches(p.margin_right_in)
        section.page_width = Inches(p.page_width_in)
        section.page_height = Inches(p.page_height_in)
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.different_first_page_header_footer = p.suppress_first_page_number
        _install_page_number(section, p.page_numbering_position)


_ALIGNMENT_FOR_POSITION = {
    "bottom-center": 1,  # WD_PARAGRAPH_ALIGNMENT.CENTER
    "bottom-right": 2,   # WD_PARAGRAPH_ALIGNMENT.RIGHT
    "bottom-left": 0,    # WD_PARAGRAPH_ALIGNMENT.LEFT
}


def _install_page_number(section, position: str) -> None:
    """Footer page number, alignment per profile.page_numbering_position."""
    footer = section.footer
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()

    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = _ALIGNMENT_FOR_POSITION.get(position, 1)

    run = p.add_run()
    _append_field(run, "PAGE \\* Arabic")


def _append_field(run, instruction: str) -> None:
    """Insert a Word field into an existing run."""
    from lxml import etree

    def el(tag: str):
        return etree.SubElement(run._element, qn(f"w:{tag}"))

    fld_begin = el("fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = etree.SubElement(run._element, qn("w:instrText"))
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "

    fld_sep = el("fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    t = etree.SubElement(run._element, qn("w:t"))
    t.text = "1"

    fld_end = el("fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
