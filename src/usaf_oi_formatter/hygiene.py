"""Whitespace, quote, and dash normalization across all runs."""

from __future__ import annotations

from docx.document import Document

from .profile import FormattingProfile, default as _default_profile

_REPLACEMENTS = [
    (".  ", ". "),
    ("  ", " "),
    ("“", '"'),
    ("”", '"'),
    ("‘", "'"),
    ("’", "'"),
]


def apply(doc: Document, profile: FormattingProfile | None = None) -> None:
    p = profile or _default_profile()
    if not p.apply_hygiene:
        return
    for p in doc.paragraphs:
        for run in p.runs:
            text = run.text
            if not text:
                continue
            for needle, replacement in _REPLACEMENTS:
                # Repeat for the double-space rule in case of chains.
                while needle in text:
                    new_text = text.replace(needle, replacement)
                    if new_text == text:
                        break
                    text = new_text
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for needle, replacement in _REPLACEMENTS:
                            run.text = run.text.replace(needle, replacement)
