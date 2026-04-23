"""Orchestrator: runs the full formatting pipeline against one .docx."""

from __future__ import annotations

import re
from pathlib import Path

import docx

from . import (
    acronyms,
    attachments,
    bullets,
    headerblock,
    hygiene,
    numbering,
    pagesetup,
    report as report_mod,
    rules,
    styles,
)
from .meta import OIMeta

_RE_NUMBERED_HEADING = re.compile(r"^(\d+(?:\.\d+){0,4})\.?\s+")


def format_file(src: Path, meta: OIMeta, output_dir: Path | None = None) -> tuple[Path, Path]:
    """Format `src` and save next to it (or into `output_dir`).

    Returns (formatted_path, report_path).
    """
    src = Path(src)
    doc = docx.Document(str(src))

    report = report_mod.ChangeReport(src)
    report.snapshot_pre(_snapshot(doc))

    report.note("page", "Applying margins and page setup")
    pagesetup.apply(doc)

    report.note("styles", "Installing OI styles")
    styles.install_or_refresh(doc)

    report.note("header", "Rebuilding DAFMAN 90-161 title block")
    headerblock.rebuild(doc, meta.with_defaults())

    report.note("walk", "Classifying paragraphs")
    _classify_paragraphs(doc)

    report.note("numbering", "Applying 1. / 1.1. / 1.1.1. numbering")
    numbering.apply(doc)

    report.note("bullets", "Normalizing bullets")
    bullets.apply(doc)

    report.note("acronyms", "Collecting acronyms")
    glossary = acronyms.collect(doc)
    report.note("acronyms", f"Found {len(glossary)} acronyms")

    report.note("attachments", "Rebuilding attachment titles")
    attachments.apply(doc, glossary)

    report.note("hygiene", "Whitespace / quotes / dashes")
    hygiene.apply(doc)

    report.diff_post(_snapshot(doc))

    out_path = _output_path(src, output_dir)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    report_path = report.write_sidecar(out_path)
    return out_path, report_path


def _output_path(src: Path, output_dir: Path | None) -> Path:
    new_name = src.stem + rules.OUTPUT_SUFFIX + src.suffix
    if output_dir is None:
        return src.with_name(new_name)
    return Path(output_dir) / new_name


def _snapshot(doc) -> dict[str, str]:
    section = doc.sections[0]
    return {
        "margin.top_emu": str(section.top_margin),
        "margin.bottom_emu": str(section.bottom_margin),
        "margin.left_emu": str(section.left_margin),
        "margin.right_emu": str(section.right_margin),
        "page.width_emu": str(section.page_width),
        "page.height_emu": str(section.page_height),
        "paragraph.count": str(len(doc.paragraphs)),
    }


# ---- paragraph classification --------------------------------------

_PROTECTED_STYLES = frozenset({
    rules.STY_TITLEBLOCK,
    rules.STY_TITLE,
    rules.STY_ATTACH_TITLE,
})


def _classify_paragraphs(doc) -> None:
    """Heuristically reassign every body paragraph to a canonical OI style.

    Skips anything the header-block builder already styled (title block,
    compliance banner, attachment titles) and anything inside a table.
    """
    for p in doc.paragraphs:
        if _is_in_table(p):
            continue
        if p.style.name in _PROTECTED_STYLES:
            continue

        text = p.text.strip()
        if not text:
            continue

        level = _leading_number_depth(text)
        if 1 <= level <= rules.MAX_NUMBER_DEPTH:
            p.style = doc.styles[rules.heading_style_for_level(level)]
        elif _is_all_caps_heading(text):
            p.style = doc.styles[rules.STY_H1]
        else:
            p.style = doc.styles[rules.STY_BODY]


def _is_in_table(paragraph) -> bool:
    from docx.oxml.ns import qn
    node = paragraph._p.getparent()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return True
        node = node.getparent()
    return False


def _leading_number_depth(text: str) -> int:
    match = _RE_NUMBERED_HEADING.match(text)
    if not match:
        return 0
    return 1 + match.group(1).count(".")


def _is_all_caps_heading(text: str) -> bool:
    if not 3 <= len(text) <= 120:
        return False
    if text != text.upper():
        return False
    return any(ch.isalpha() for ch in text)
