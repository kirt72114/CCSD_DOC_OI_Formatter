"""End-to-end: build a throwaway template + draft and apply the template."""

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

from usaf_oi_formatter import formatter, rules


def _make_draft(path: Path) -> None:
    doc = docx.Document()
    doc.add_paragraph("DRAFT CONTENT GOES HERE")
    doc.add_paragraph("Some body text with a  double space and a  smart quote: “foo”.")
    doc.add_paragraph("1. Scope.")
    doc.add_paragraph("Applies to all personnel.")
    doc.save(str(path))


def _make_template(path: Path) -> None:
    doc = docx.Document()
    p = doc.add_paragraph("Template body placeholder")
    # Tweak page margins and give it a distinctive style so we can detect
    # that the template copy worked.
    from docx.shared import Inches
    s = doc.sections[0]
    s.top_margin = Inches(1.5)
    s.bottom_margin = Inches(1.5)
    s.left_margin = Inches(1.25)
    s.right_margin = Inches(1.25)
    doc.save(str(path))


def test_no_template_runs_only_hygiene(tmp_path: Path) -> None:
    src = tmp_path / "draft.docx"
    _make_draft(src)
    out, report = formatter.format_file(src)
    assert out.exists()
    assert out.name == f"draft{rules.OUTPUT_SUFFIX}.docx"
    # Hygiene should have collapsed the double space.
    result = docx.Document(str(out))
    text = "\n".join(p.text for p in result.paragraphs)
    assert "double space" in text
    assert "  " not in text  # collapsed
    # Smart quote converted to straight.
    assert '"foo"' in text
    assert report.read_text(encoding="utf-8")


def test_template_applies_margins(tmp_path: Path) -> None:
    draft = tmp_path / "draft.docx"
    template = tmp_path / "template.docx"
    _make_draft(draft)
    _make_template(template)

    out, _ = formatter.format_file(draft, template=template)
    assert out.exists()

    from docx.shared import Inches
    result = docx.Document(str(out))
    section = result.sections[0]
    # The draft was built with the 1" default; the template used 1.5"/1.25".
    # Expect the template's margins on the output.
    assert section.top_margin == Inches(1.5)
    assert section.left_margin == Inches(1.25)
